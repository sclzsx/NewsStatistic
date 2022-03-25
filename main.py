# 以一个社区为单位，关键词出现的新闻篇数占该社区总新闻篇数的比例

import numpy as np
import pandas as pd
import docx
import sys


def extract_keywords_from_docx():
    file = docx.Document("keywords2.0.docx")

    keywords = []  # 国家，公司，党，其他
    k = 0
    tmp = []
    pass_words = ['关键词', '国家级', '省级', '市级', '区级']
    for i in range(len(file.paragraphs)):
        text = file.paragraphs[i].text
        if len(text) > 0:
            pass_flag = 0
            for pass_word in pass_words:
                if pass_word in text:
                    pass_flag = 1
                    break
            if not pass_flag:
                tmp.append(text)

        if len(text) == 0:
            k = k + 1
            keywords.append(tmp)
            tmp = []

    return keywords


keywords = extract_keywords_from_docx()
print('完成docx读取')
print(keywords[0])
print(keywords[1])
print(keywords[2])
print(keywords[3])
print()

data = np.array(pd.read_excel('community_news3.0.xlsx', sheet_name='社区新闻原始数据'))
results = np.array(pd.read_excel('community_news3.0.xlsx', sheet_name='结果'))

d_communitiy_names = data[:, 2]
d_communitiy_news = data[:, 5]
r_communitiy_names = results[:, 1]
r_communitiy_newsnums = results[:, 2]

print('完成excel读取')
print(r_communitiy_names)
print(results.shape, data.shape)
print()


def check_keywords_in_news(keywords, news):
    for keyword in keywords:
        if keyword in news:
            return 1
    return 0


count = np.zeros((r_communitiy_names.shape[0], 17))
for i, r_communitiy_name in enumerate(r_communitiy_names):
    total_num = r_communitiy_newsnums[i]
    country_num, company_num, party_num, others_num = 0, 0, 0, 0
    country_party_num, country_company_num, party_company_num, country_party_company_num = 0, 0, 0, 0
    for j, name in enumerate(d_communitiy_names):
        if name == r_communitiy_name:
            news = d_communitiy_news[j]
            if not isinstance(news, str):
                continue
            country_num = country_num + check_keywords_in_news(keywords[0], news)
            company_num = company_num + check_keywords_in_news(keywords[1], news)
            party_num = party_num + check_keywords_in_news(keywords[2], news)
            others_num = others_num + check_keywords_in_news(keywords[3], news)

            country_party_num = country_party_num + (check_keywords_in_news(keywords[0], news) & check_keywords_in_news(keywords[2], news))
            country_company_num = country_company_num + (check_keywords_in_news(keywords[0], news) & check_keywords_in_news(keywords[1], news))
            party_company_num = party_company_num + (check_keywords_in_news(keywords[2], news) & check_keywords_in_news(keywords[1], news))
            country_party_company_num = country_party_company_num + (check_keywords_in_news(keywords[0], news) & check_keywords_in_news(keywords[1], news) & check_keywords_in_news(keywords[2], news))

    print('name:{}, total_num:{}, country_num:{}, company_num:{}, party_num:{}, others_num:{}'.format(r_communitiy_name,
                                                                                                      total_num,
                                                                                                      country_num,
                                                                                                      company_num,
                                                                                                      party_num,
                                                                                                      others_num))
    count[i, 0] = total_num
    count[i, 1] = country_num
    count[i, 2] = country_num / total_num
    count[i, 3] = company_num
    count[i, 4] = company_num / total_num
    count[i, 5] = party_num
    count[i, 6] = party_num / total_num
    count[i, 7] = others_num
    count[i, 8] = others_num / total_num

    count[i, 9]  = country_party_num
    count[i, 10] = country_party_num / total_num
    count[i, 11] = country_company_num
    count[i, 12] = country_company_num / total_num
    count[i, 13] = party_company_num
    count[i, 14] = party_company_num / total_num
    count[i, 15] = country_party_company_num
    count[i, 16] = country_party_company_num / total_num
print('完成统计')
print()

pd.DataFrame(count).to_csv('count.csv', index=False, header=False)
count_read = np.array(pd.read_csv('count.csv', header=None))
print('完成结果保存')
